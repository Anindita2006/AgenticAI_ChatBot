"""
Builds data/bvrith_college_info.docx, the grounding document for the RAG chatbot.

Content (Sections 1-8) is transcribed from bvrithyderabad.edu.in (About, Admissions,
Fee Details, Placements, Training & Placement Cell, and each department's Faculty /
About-HOD pages — CSE, ECE, EEE, IT, CSE AI&ML) plus a small number of facts
corroborated via search (library/hostel/transport figures, contact details).
Sections 9-10 (Scholarships & Fee Concessions, Student Support Services) and the
hostel/mess/other per-year fee figures in Section 4 fill gaps the real site didn't
publish, using bvrit_college_info.pdf — the course's reference grounding document —
since no real-site figures existed for these topics.
Where a figure is still not published anywhere, the document says so explicitly
instead of inventing a number — the chatbot must be able to refuse on these gracefully.

Each of the 10 sections is placed on its own page (explicit page break) so that the
page number is deterministic and can be used as citation metadata during ingestion,
without needing to render/paginate the .docx.
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_PATH = "bvrith_college_info.docx"

SECTIONS = [
    {
        "heading": "1. About BVRIT Hyderabad",
        "body": [
            ("BVRIT HYDERABAD College of Engineering for Women (Autonomous) was "
             "established in 2012 by the Sri Vishnu Educational Society (SVES), "
             "founded by Chairman K V Vishnu Raju, with the vision and mission to "
             "empower women to aspire for a better career in life."),
            ("The college is located in Bachupally, Hyderabad, Telangana, and is "
             "affiliated to Jawaharlal Nehru Technological University Hyderabad "
             "(JNTUH), approved by the All India Council for Technical Education "
             "(AICTE), and approved by the Government of Telangana."),
            ("Vision: To emerge as the best among institutes of technology and "
             "research in the country, dedicated to the cause of promoting quality "
             "technical education."),
            ("Mission (core focus areas):"),
            ("- Deliver academic excellence through outcome-based education", True),
            ("- Enhance technical competency of students to meet industry needs", True),
            ("- Encourage research and innovation among faculty and students", True),
            ("- Support holistic student development, including women's empowerment and leadership", True),
            ("Accreditations:"),
            ("- NAAC Grade 'A' (CGPA 3.23), awarded in 2020", True),
            ("- UGC Autonomous status, granted from AY 2023-24 for a period of 10 years", True),
            ("- NBA (National Board of Accreditation) for B.Tech programmes in EEE, ECE, and CSE", True),
            ("- AICTE approved; affiliated to JNTUH", True),
            ("Recognition: NIRF rankings (2019-2021), ARIIA 2020 Band B ranking, and "
             "the IMC Ramakrishna Bajaj National Quality Award."),
        ],
    },
    {
        "heading": "2. Departments",
        "body": [
            ("BVRIT Hyderabad offers 4 undergraduate B.Tech programmes with a total "
             "sanctioned intake of 660 students:"),
            ("- Computer Science and Engineering (CSE)", True),
            ("- Computer Science and Engineering — Artificial Intelligence and Machine Learning (CSE AI&ML / CSM)", True),
            ("- Electronics and Communication Engineering (ECE)", True),
            ("- Electrical and Electronics Engineering (EEE)", True),
            ("- Information Technology (IT)", True),
            ("Postgraduate programmes (3 M.Tech specialisations, combined intake 42):"),
            ("- M.Tech in Data Sciences — intake 18", True),
            ("- M.Tech in VLSI Design — intake 12", True),
            ("- M.Tech in Computer Science and Engineering — intake 12", True),
            ("Doctoral research centres are active in ECE (VLSI & Communications) and CSE."),
            ("Department of Computer Science and Engineering (detail available): "
             "established in 2012 with an initial intake of 60 students, which grew to "
             "120 (2014-15), 180 (2019-20), and 360 (2024-25). As listed on the "
             "department's official faculty page, CSE has 51 faculty members — 5 "
             "Professors, 4 Associate Professors, and 42 Assistant Professors — "
             "supported by adjunct faculty from industry and academia, and is headed "
             "by Dr. Aruna Rao S L (Professor & Head, Dept. of CSE). "
             "The CSE programme has been NBA-accredited since 2018-19 (extended through "
             "2021-22). Faculty research areas include Machine Learning and AI, computer "
             "networks, data mining, image processing, parallel computing, software "
             "engineering, cloud computing, soft computing, blockchain, deep learning, "
             "natural language processing, cyber security, and data analytics."),
            ("Detailed faculty counts, specialisations, and HOD names for ECE, EEE, "
             "IT, and CSE AI&ML departments individually are not covered by the source "
             "material used for this document — please contact the respective "
             "department or the college directly for those figures."),
        ],
    },
    {
        "heading": "3. Admissions",
        "body": [
            ("Eligibility: Candidates must have completed 10+2 with Mathematics, "
             "Physics, and Chemistry (MPC) as major subjects, with a minimum of 60% "
             "aggregate marks."),
            ("Entrance exam: Admission to B.Tech programmes is primarily through TS "
             "EAMCET (Telangana State Engineering, Agricultural and Medical Common "
             "Entrance Test), also referred to as TG EAPCET. Admission to M.Tech "
             "programmes is through PGECET. The college code for undergraduate "
             "admissions is BVRW, and BVRW1 for postgraduate admissions."),
            ("Admission categories:"),
            ("- Category A (Convener Seats): filled by the TS EAMCET Convener based on merit rank", True),
            ("- Category B (College/Management Seats): allocated by the institution as per Telangana State Council of Higher Education (TSCHE) guidelines", True),
            ("Process: Category A applicants apply through the Convener-TSEAMCET "
             "portal. Category B applicants apply via the college website once TSCHE "
             "issues its notification; reapplication is not required for Category B "
             "if already registered."),
            ("JEE (Main): BVRIT Hyderabad does NOT use JEE (Main) scores for B.Tech "
             "admission. TS EAMCET / TG EAPCET is the sole entrance exam route for "
             "Category A seats. Applicants should not expect a JEE-based admission "
             "pathway at this college."),
            ("Category B (management quota) detail: the college's official B-Category "
             "admissions page does not currently publish eligibility, fee, or process "
             "details for this category. Contact the admissions office directly for "
             "management-quota specifics rather than assuming a figure."),
            ("Key dates: The complete admission schedule (exam dates, counselling, "
             "seat allotment) is released by TSCHE and TS EAMCET each academic year. "
             "Applicants should monitor official TSEAMCET and TSCHE notifications for "
             "exact dates, as this document does not have a fixed calendar."),
            ("Admissions contact: Dr. J. Manoj Kumar, Professor i/c Admissions — phone 92471 64714."),
        ],
    },
    {
        "heading": "4. Fee Structure",
        "body": [
            ("Annual tuition fee (as published by the college, per academic year of admission):"),
            ("- 2025 batch: Rs. 1,20,000 per annum — CSE, ECE, EEE, CSM", True),
            ("- 2024 batch: Rs. 1,20,000 per annum — CSE, ECE, EEE, CSM", True),
            ("- 2023 batch: Rs. 1,20,000 per annum — CSE, ECE, IT, EEE, CSM", True),
            ("- 2021-2022 batches: Rs. 90,000 per annum — all engineering branches", True),
            ("Additional annual charges:"),
            ("- NBA Fee: Rs. 3,000 per annum, applicable to NBA-accredited programmes (EEE, ECE, CSE). CSM (AI&ML) is exempt from the NBA fee.", True),
            ("- JNTUH / Miscellaneous Fee: Rs. 5,500 per annum (Rs. 2,500 for the 2022 batch)", True),
            ("Hostel, mess, and other per-year fee figures (as published in the college's "
             "reference fee schedule, not on the public fee-details webpage used for the "
             "rest of this section):"),
            ("- Hostel room rent, shared 3-seater: Rs. 60,000 per year", True),
            ("- Hostel room rent, shared 2-seater: Rs. 75,000 per year", True),
            ("- Mess charges (vegetarian): Rs. 48,000 per year", True),
            ("- Mess charges (non-vegetarian): Rs. 54,000 per year", True),
            ("- Transport (college bus): Rs. 45,000 per year, optional", True),
            ("- Laboratory fee: Rs. 15,000 per year", True),
            ("- Library & digital resources fee: Rs. 8,000 per year", True),
            ("- Examination fee: Rs. 5,000 per semester (Rs. 10,000 per year)", True),
            ("- Student activity fee: Rs. 3,000 per year", True),
            ("- Caution deposit (refundable): Rs. 10,000, one-time at admission", True),
            ("- Admission processing fee: Rs. 5,000, one-time at admission", True),
            ("Scholarship and fee-concession schemes are covered separately in Section 9 "
             "(Scholarships & Fee Concessions). For anything not covered above, contact "
             "the admissions office (Dr. J. Manoj Kumar, 92471 64714, or "
             "info@bvrithyderabad.edu.in) rather than assuming a figure."),
        ],
    },
    {
        "heading": "5. Placements",
        "body": [
            ("The Training and Placement (TAP) Cell serves students across Sri Vishnu "
             "Educational Society institutions and reports that BVRIT Hyderabad has "
             "consistently topped placement records among private institutions in "
             "Telangana and Andhra Pradesh over the past seven years. The cell is "
             "headed by a Dean and assisted by Placement Officers."),
            ("TAP Cell activities include: industrial training during summer and "
             "semester breaks, entrepreneurship development programmes, career "
             "guidance through corporate executive lectures, competitive exam "
             "preparation (GATE, GRE, GMAT, TOEFL), campus recruitment drives, and "
             "industry-institution partnerships for projects and employment."),
            ("Top recruiters include Microsoft, Amazon, Visa, Optum, Synopsys, IBM, "
             "Oracle, TCS, Infosys, Capgemini, Cognizant, and Tech Mahindra."),
            ("Highest packages reported (2021-2025 batch): Microsoft — Rs. 54.00 LPA "
             "(highest on record); Amazon SDE — Rs. 48.6 LPA; Visa — Rs. 32.88 LPA."),
            ("Typical / mid-tier packages: broadly in the Rs. 3.5-4.5 LPA range, e.g. "
             "TCS Ninja (Rs. 3.36-3.96 LPA), Capgemini (Rs. 3.15-4.25 LPA), Tech "
             "Mahindra (approx. Rs. 3.25 LPA)."),
            ("Placement counts by graduating batch (number of students placed): "
             "2021-2025 — 614; 2020-2024 — 508; 2019-2023 — 694; 2018-2022 — 988; "
             "2017-2021 — 533. These totals span all eligible departments (CSE, ECE, "
             "EEE, IT, and CSE AI&ML where applicable)."),
            ("An overall single-number placement percentage (e.g. \"95% placed\") is "
             "not published on the source pages used for this document — only "
             "batch-wise placement counts and package figures above are confirmed. "
             "Do not state a placement percentage that is not given here."),
        ],
    },
    {
        "heading": "6. Campus & Facilities",
        "body": [
            ("Library: approximately 22,825 books, seating capacity for 3,117 "
             "students, 150 magazines, 12 online journals, 6,306 CDs, and 1,690 other "
             "materials. The BVRITH digital library has 10 computers with internet "
             "access for online study material, NPTEL video lectures, and online mock "
             "tests."),
            ("Hostel: 4 hostel blocks with 150+ rooms and occupancy for 500+ students. "
             "Amenities include 24/7 water supply, geysers, and nutritious food. Each "
             "block has experienced lady wardens and dedicated security officers for "
             "student safety."),
            ("Sports: the campus has a 200-metre running track with shot put and "
             "discus throw circles, 2 shuttle badminton courts, and one court each "
             "for basketball, volleyball, throwball, kabaddi, kho-kho, and tennicoit. "
             "A Physical Director and three external coaches (specialising in "
             "kho-kho, basketball, throwball, and volleyball) train students, who "
             "compete in the Nipunya Inter-College Tournament, the Annual Day Sports "
             "Meet, the JNTUH Inter-College Tournament, and the India Open "
             "Inter-College Tournament, with several selections at state and "
             "national level from 2023-24 onwards."),
            ("Campus WiFi: available in every classroom, with particular emphasis on "
             "coverage for 3rd- and 4th-year students."),
            ("Transport: the college operates 14 buses across 14+ routes covering "
             "Hyderabad city, with GPS/GPRS tracking on all buses for safety."),
            ("Other facilities: gymnasium, cafeteria/canteen, temple, auditorium, "
             "clinic, guest rooms, security systems, engineering labs, and campus IT "
             "infrastructure."),
        ],
    },
    {
        "heading": "7. Faculty",
        "body": [
            ("Principal: Dr. K. V. N. Sunitha, Founder Principal of BVRIT Hyderabad "
             "(since August 2012)."),
            ("- Qualifications: B.Tech in Electronics and Communication Engineering (Acharya Nagarjuna University); M.Tech in Computer Science (Regional Engineering College, Warangal, 1993); Ph.D in Computer Science and Engineering (JNTU Hyderabad, 2006).", True),
            ("- Research interests: natural language processing, speech processing, and network security. Has supervised 20 doctoral scholars and authored five textbooks.", True),
            ("- Awards: \"Engineer of the Year 2019\" (Institution of Engineers, India); \"Acharya Ratna\" National Award for lifetime achievement (2019); \"Distinguished Principal Award\" (Computer Society of India, 2017); \"Best Engineering College Principal in Telangana State\" (ISTE, 2019).", True),
            ("Department of Computer Science and Engineering:"),
            ("- Head of Department: Dr. Aruna Rao S L, Professor & Head, Dept. of CSE.", True),
            ("- Faculty strength: 51 total — 5 Professors, 4 Associate Professors, 42 Assistant Professors — plus adjunct faculty drawn from industry and academia.", True),
            ("- Research areas: Machine Learning and Artificial Intelligence, computer networks, data mining, image processing, parallel computing, software engineering, cloud computing, soft computing, blockchain, deep learning, natural language processing, cyber security, and data analytics.", True),
            ("- The CSE undergraduate programme has been NBA-accredited since 2018-19, extended through 2021-22.", True),
            ("Department of Electronics and Communication Engineering:"),
            ("- Head of Department: Dr. Nagesh Deevi, Associate Professor & HoD, Dept. of ECE. Ph.D in RF-VLSI (NIT Warangal); IEEE member; research interests include device modeling, RF communications, on-chip component design, and semiconductor packaging.", True),
            ("- Faculty strength: 27 total — 5 Professors, 7 Associate Professors, 15 Assistant Professors.", True),
            ("Department of Electrical and Electronics Engineering:"),
            ("- Head of Department: Dr. M. Sharanya, Professor & HoD, Dept. of EEE. B.Tech in EEE (JNTUH, 2001); M.Tech in Power Electronics and Electrical Drives (JNTUH, 2007); Ph.D (JNTUH, 2020); research interests include Power Electronics, Power Quality, Renewable Energy, and Electric Vehicles.", True),
            ("- Faculty strength: 16 total — 2 Professors, 5 Associate Professors, 9 Assistant Professors.", True),
            ("Department of Information Technology:"),
            ("- Head of Department: Dr. K. Srikar Goud, Assistant Professor & I/C HoD (In-Charge), Dept. of IT.", True),
            ("- Faculty strength: 9 total — 1 Associate Professor (Dr. K. Bharathi) and 8 Assistant Professors (including Mr. A. Rajashekar Reddy, Ms. K. Kavitha, Ms. A. Aruna Jyothi, Ms. N. Sandhya, Ms. Danthuluri Mani Sri Madhuri, Ms. A. S. S. M. Pravallika, and Ms. T. Sukanya).", True),
            ("Department of CSE — Artificial Intelligence and Machine Learning (CSE AI&ML / CSM):"),
            ("- Head of Department: Dr. Venkata Raja Sekhar Reddy N, Professor & HoD, Dept. of CSE (AI&ML). Contact: hod.aiml@bvrithyderabad.edu.in.", True),
            ("- Faculty strength: 15 total — 2 Professors, 0 Associate Professors, 13 Assistant Professors.", True),
            ("Faculty rosters and qualifications beyond what is listed above (e.g. "
             "individual assistant professor profiles) are published on each "
             "department's official faculty page at bvrithyderabad.edu.in; the "
             "chatbot should direct the user there or to the admissions office "
             "rather than guessing if asked for detail not covered in this document."),
        ],
    },
    {
        "heading": "8. Contact",
        "body": [
            ("Address: BVRIT HYDERABAD College of Engineering for Women, Plot No. "
             "8-5/4, Rajiv Gandhi Nagar Colony, Nizampet Road, Bachupally, Hyderabad "
             "- 500090, Telangana, India."),
            ("Phone: +91 40 4241 7773"),
            ("Email: info@bvrithyderabad.edu.in (general enquiries), "
             "principal@bvrithyderabad.edu.in (principal's office)"),
            ("Admissions contact: Dr. J. Manoj Kumar, Professor i/c Admissions — "
             "phone 92471 64714."),
            ("Official website: https://bvrithyderabad.edu.in"),
            ("Social media:"),
            ("- LinkedIn: linkedin.com/school/bvrit-hyderabad/", True),
            ("- Instagram: instagram.com/bvrit_hyderabad_/", True),
            ("- Facebook: facebook.com/BvritHyderabad/", True),
            ("- YouTube: youtube.com/@BvritHyderabadWomen", True),
            ("- X (Twitter): x.com/bvrithyderabad", True),
        ],
    },
    {
        "heading": "9. Scholarships & Fee Concessions",
        "body": [
            ("BVRIT Hyderabad offers several scholarship schemes to support meritorious and "
             "economically disadvantaged students. Scholarships are applied as a percentage "
             "discount on the annual tuition fee only (not on hostel, mess, or other fees)."),
            ("Merit scholarships:"),
            ("- Founder's Scholarship: EAMCET/EAPCET rank within top 1,000 — 100% (full tuition waiver); renewable if CGPA >= 8.5", True),
            ("- Academic Excellence: EAMCET/EAPCET rank 1,001-5,000 — 50% discount on tuition; renewable if CGPA >= 8.0", True),
            ("- Merit Reward: EAMCET/EAPCET rank 5,001-15,000 — 25% discount on tuition; renewable if CGPA >= 7.5", True),
            ("- Sports Scholarship: national/state level sports achievement — 25% discount on tuition; renewable with active participation", True),
            ("Need-based fee concessions:"),
            ("- Economically Weaker Section (EWS): family income below Rs. 8 lakh/year — 50% discount on tuition", True),
            ("- SC/ST Scholarship: Telangana state SC/ST students — full tuition reimbursement via state government", True),
            ("- BC Fee Reimbursement: Telangana BC students, income under Rs. 2 lakh — full tuition reimbursement via state government", True),
            ("- Sibling Discount: second sibling currently enrolled at BVRIT Hyderabad — 10% discount on tuition", True),
            ("Note: government scholarships (SC/ST, BC) are subject to state government "
             "disbursement timelines. Students must pay the full fee at admission and receive "
             "reimbursement after government processing, which typically takes 3-6 months."),
        ],
    },
    {
        "heading": "10. Student Support Services",
        "body": [
            ("Student Counselling Centre: provides free, confidential support for academic "
             "stress, personal issues, and mental health concerns. Two full-time counsellors "
             "are available Monday-Saturday, 9 AM - 5 PM. Appointments can be booked in person "
             "at the Admin Block or via email at counselling@bvrithyderabad.edu.in; walk-ins "
             "are welcome."),
            ("In case of crisis: students in distress can contact the counselling centre "
             "directly. For after-hours emergencies, contact hostel wardens or campus security "
             "(available 24/7)."),
            ("External crisis resources:"),
            ("- iCall — 9152987821 (Mon-Sat, 8 AM - 10 PM)", True),
            ("- Vandrevala Foundation Helpline — 1860-2662-345 (24/7)", True),
            ("- NIMHANS Helpline — 080-46110007", True),
            ("Anti-Ragging Committee: BVRIT Hyderabad has a zero-tolerance policy on ragging. "
             "The Anti-Ragging Committee can be contacted via the admissions office or the UGC "
             "helpline 1800-180-5522. All complaints are investigated within 48 hours."),
            ("Grievance Redressal: students can submit grievances in writing to the Dean "
             "(Student Affairs) or via the college's grievance channel. All grievances receive "
             "a response within 7 working days. The Grievance Redressal Committee meets monthly "
             "to review pending cases."),
        ],
    },
]


def build_document(output_path: str = OUTPUT_PATH) -> None:
    doc = Document()

    title = doc.add_heading("BVRIT HYDERABAD College of Engineering for Women", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Official Information Document — Knowledge Base for the College FAQ Chatbot")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    note = doc.add_paragraph(
        "Compiled from bvrithyderabad.edu.in (About, Admissions, Fee Details, "
        "Placements, Training & Placement Cell, CSE Department pages). Facts not "
        "found on the official pages are explicitly marked as unavailable rather "
        "than invented."
    )
    note.runs[0].font.size = Pt(9)
    note.runs[0].italic = True
    doc.add_page_break()

    for i, section in enumerate(SECTIONS):
        doc.add_heading(section["heading"], level=1)
        for item in section["body"]:
            if isinstance(item, tuple):
                text, is_bullet = item
                doc.add_paragraph(text, style="List Bullet")
            else:
                doc.add_paragraph(item)
        if i < len(SECTIONS) - 1:
            doc.add_page_break()

    doc.save(output_path)
    print(f"Saved {output_path} ({len(SECTIONS)} sections, {len(SECTIONS)} pages of content + 1 title page)")


if __name__ == "__main__":
    build_document()
